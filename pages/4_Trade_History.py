"""
Trade History Page
View all trades with filtering, deletion, and expandable details
"""

import streamlit as st
import pandas as pd
import io


from utils.paths import get_database_path, get_trade_screenshot_dir
from database.trades_db import get_filtered_trades, get_trade_screenshots, delete_trade, get_all_trades, update_trade_open_fields
from database.accounts_db import get_all_accounts
from database.strategies_db import get_all_strategies
from components.account_selector import render_account_selector
from utils.formatters import (
    format_currency,
    format_percentage,
    format_r_multiple,
    format_direction,
    format_duration
)
from datetime import datetime
from pathlib import Path
from PIL import Image
import logging
from utils.cache_manager import clear_cache_after_trade_operation
from utils.png_icons import *

logger = logging.getLogger(__name__)


# ============================================================================
# HELPER FUNCTIONS FOR ORGANIZED EXPORTS
# ============================================================================

def get_exports_base_dir(db_path):
    """Get base exports directory"""
    base_dir = Path(db_path).parent / "Exports"
    base_dir.mkdir(parents=True, exist_ok=True)
    return base_dir


def get_csv_exports_dir(db_path):
    """Get CSV exports directory"""
    base_dir = get_exports_base_dir(db_path)
    csv_dir = base_dir / "CSV_Trades_Exports"
    csv_dir.mkdir(parents=True, exist_ok=True)
    return csv_dir


def get_account_export_dir(db_path, account_name):
    """Get or create account-specific export directory"""
    base_dir = get_exports_base_dir(db_path)
    account_folder = account_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    account_dir = base_dir / account_folder
    account_dir.mkdir(parents=True, exist_ok=True)
    return account_dir


def get_week_folder_name(trade_date_str):
    """
    Generate week folder name from trade date
    Format: YYYY_MMM_Wxx (e.g., 2026_Feb_W03)
    Weeks start on Monday (ISO calendar)
    """
    try:
        if isinstance(trade_date_str, str):
            dt = datetime.strptime(trade_date_str, "%Y-%m-%d")
        else:
            dt = trade_date_str

        year = dt.year
        month_abbr = dt.strftime("%b")  # Jan, Feb, Mar, etc.
        week_of_year = dt.isocalendar()[1]  # ISO week number (weeks start Monday)

        return f"{year}_{month_abbr}_W{week_of_year:02d}"
    except:
        return "Unknown_Week"


def get_organized_export_path(db_path, account_name, trade_date_str):
    """
    Get organized export path with account/week structure
    Format: Exports/Account_Name/2026_Feb_W03/

    Args:
        db_path: Database path
        account_name: Account name
        trade_date_str: Trade date in YYYY-MM-DD format

    Returns:
        Path object for the organized directory
    """
    account_dir = get_account_export_dir(db_path, account_name)
    week_folder = get_week_folder_name(trade_date_str)
    organized_dir = account_dir / week_folder
    organized_dir.mkdir(parents=True, exist_ok=True)

    return organized_dir


# ============================================================================
# DOCX EXPORT FUNCTION (Modified for week-only organization)
# ============================================================================

def export_trade_to_docx(trade, account_name, all_strategies, screenshots, db_path):
    """
    Export individual trade details to DOCX (Word document) with screenshots
    Organized by: Exports/Account_Name/YYYY_MMM_Wxx/
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image

    doc = Document()

    # Title
    title = doc.add_heading(f"Trade Report - Trade #{trade['trade_id']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Account info
    doc.add_paragraph(f"Account: {account_name}")
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # Trade Details Section
    doc.add_heading("Trade Details", level=1)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = "Trade ID"
    table.rows[0].cells[1].text = f"#{trade['trade_id']}"

    table.rows[1].cells[0].text = "Symbol"
    table.rows[1].cells[1].text = str(trade['symbol'])

    table.rows[2].cells[0].text = "Direction"
    table.rows[2].cells[1].text = str(trade['direction'])

    table.rows[3].cells[0].text = "Status"
    table.rows[3].cells[1].text = str(trade['status'])

    table.rows[4].cells[0].text = "Entry Date"
    table.rows[4].cells[1].text = f"{trade['entry_date']} {trade['entry_time']}"

    table.rows[5].cells[0].text = "Entry Price"
    table.rows[5].cells[1].text = f"${trade['entry_price']:.2f}"

    if trade.get('multi_entry_price'):
        table.rows[6].cells[0].text = "Weighted Avg Entry"
        table.rows[6].cells[1].text = f"${trade['multi_entry_price']:.2f}"

    doc.add_paragraph()

    # Trade Setup Section
    doc.add_heading("Trade Setup", level=1)

    setup_table = doc.add_table(rows=3, cols=2)
    setup_table.style = 'Light Grid Accent 1'

    trading_env = trade.get('trading_environment') or 'N/A'
    setup_table.rows[0].cells[0].text = "Trading Environment"
    setup_table.rows[0].cells[1].text = str(trading_env)

    strategy_name = "No Strategy"
    if trade.get('strategy_id'):
        for s in all_strategies:
            if s['strategy_id'] == trade['strategy_id']:
                strategy_name = s['strategy_name']
                break

    setup_table.rows[1].cells[0].text = "Strategy"
    setup_table.rows[1].cells[1].text = str(strategy_name)

    trigger = trade.get('trigger') or 'N/A'
    setup_table.rows[2].cells[0].text = "Trigger"
    setup_table.rows[2].cells[1].text = str(trigger)

    doc.add_paragraph()

    # Position Sizing & Risk Section
    doc.add_heading("Position Sizing & Risk", level=1)

    risk_table = doc.add_table(rows=7, cols=2)
    risk_table.style = 'Light Grid Accent 1'

    risk_table.rows[0].cells[0].text = "Position Size"
    risk_table.rows[0].cells[1].text = f"{trade['position_size']:.8f} units"

    risk_table.rows[1].cells[0].text = "Stop Loss"
    risk_table.rows[1].cells[1].text = f"${trade['stop_loss']:.2f}"

    risk_table.rows[2].cells[0].text = "Target"
    risk_table.rows[2].cells[1].text = f"${trade.get('target', 0):.2f}" if trade.get('target') else "N/A"

    risk_table.rows[3].cells[0].text = "Risk Amount"
    risk_table.rows[3].cells[1].text = f"${trade['risk_amount']:.2f}"

    risk_table.rows[4].cells[0].text = "Risk %"
    risk_table.rows[4].cells[1].text = f"{trade['risk_percentage']:.2f}%"

    risk_table.rows[5].cells[0].text = "FTA"
    risk_table.rows[5].cells[1].text = f"${trade.get('fta', 0):.2f}" if trade.get('fta') else "N/A"

    risk_table.rows[6].cells[0].text = "Prospective R"
    risk_table.rows[6].cells[1].text = f"{trade.get('prospective_r', 0):.2f}R" if trade.get('prospective_r') else "N/A"

    doc.add_paragraph()

    # Trade Results (if closed)
    if trade['status'] == 'CLOSED':
        doc.add_heading("Trade Results", level=1)

        results_table = doc.add_table(rows=9, cols=2)
        results_table.style = 'Light Grid Accent 1'

        results_table.rows[0].cells[0].text = "Exit Date"
        results_table.rows[0].cells[1].text = f"{trade.get('exit_date', 'N/A')} {trade.get('exit_time', '')}"

        results_table.rows[1].cells[0].text = "Exit Price"
        results_table.rows[1].cells[1].text = f"${trade.get('multi_exit_price', 0):.2f}" if trade.get('multi_exit_price') else "N/A"

        results_table.rows[2].cells[0].text = "Target Hit"
        results_table.rows[2].cells[1].text = "Yes" if trade.get('target_hit') else "No"

        results_table.rows[3].cells[0].text = "Stop Loss Hit"
        results_table.rows[3].cells[1].text = "Yes" if trade.get('stop_loss_hit') else "No"

        results_table.rows[4].cells[0].text = "MAE"
        results_table.rows[4].cells[1].text = f"${trade.get('mae', 0):.2f}" if trade.get('mae') else "N/A"

        results_table.rows[5].cells[0].text = "MFE"
        results_table.rows[5].cells[1].text = f"${trade.get('mfe', 0):.2f}" if trade.get('mfe') else "N/A"

        results_table.rows[6].cells[0].text = "Duration"
        results_table.rows[6].cells[1].text = str(trade.get('trade_duration', 'N/A'))

        results_table.rows[7].cells[0].text = "Total R"
        results_table.rows[7].cells[1].text = f"{trade.get('total_r', 0):.2f}R" if trade.get('total_r') else "N/A"

        trade_pnl = trade.get('pnl', 0) or 0.0
        results_table.rows[8].cells[0].text = "P&L"
        results_table.rows[8].cells[1].text = f"${trade_pnl:.2f}"

        doc.add_paragraph()

        # Trade Grades
        doc.add_heading("Trade Grades", level=1)

        grades_table = doc.add_table(rows=2, cols=2)
        grades_table.style = 'Light Grid Accent 1'

        grade_mental = trade.get('grade_mentally') or 'Not Graded'
        grades_table.rows[0].cells[0].text = "Grade Mentally"
        grades_table.rows[0].cells[1].text = str(grade_mental)

        grade_tech = trade.get('grade_technically') or 'Not Graded'
        grades_table.rows[1].cells[0].text = "Grade Technically"
        grades_table.rows[1].cells[1].text = str(grade_tech)

        doc.add_paragraph()

    # Trade Notes
    doc.add_heading("Trade Notes", level=1)

    if trade.get('setup_idea'):
        doc.add_heading("Setup Idea", level=2)
        doc.add_paragraph(str(trade['setup_idea']))

    if trade.get('trade_notes_entry'):
        doc.add_heading("Trade Notes (Entry)", level=2)
        doc.add_paragraph(str(trade['trade_notes_entry']))

    if trade.get('trade_notes_management'):
        doc.add_heading("Trade Notes (Management)", level=2)
        doc.add_paragraph(str(trade['trade_notes_management']))

    if trade.get('trade_notes_closing'):
        doc.add_heading("Trade Notes (Closing)", level=2)
        doc.add_paragraph(str(trade['trade_notes_closing']))

    if trade.get('reason_for_closing'):
        doc.add_heading("Reason for Closing", level=2)
        doc.add_paragraph(str(trade['reason_for_closing']))

    if trade.get('final_notes'):
        doc.add_heading("Final Notes", level=2)
        doc.add_paragraph(str(trade['final_notes']))

    # Screenshots section - WITH ACTUAL IMAGES
    if screenshots and len(screenshots) > 0:
        doc.add_page_break()
        doc.add_heading("Screenshots", level=1)

        for idx, screenshot in enumerate(screenshots, 1):
            screenshot_type = screenshot.get('screenshot_type', 'OTHER')
            file_path = screenshot.get('file_path', '')

            if file_path:
                screenshot_path = Path(file_path)

                if screenshot_path.exists():
                    try:
                        # Add screenshot type and name
                        doc.add_heading(f"{screenshot_type} - {screenshot_path.name}", level=2)

                        # Open image to get dimensions
                        img = Image.open(screenshot_path)
                        img_width, img_height = img.size

                        # Calculate aspect ratio
                        aspect_ratio = img_height / img_width

                        # Set max width (6 inches for Word document)
                        max_width = 6.0
                        doc_width = max_width
                        doc_height = doc_width * aspect_ratio

                        # If height is too large, scale down
                        if doc_height > 8.0:
                            doc_height = 8.0
                            doc_width = doc_height / aspect_ratio

                        # Add image to document
                        doc.add_picture(str(screenshot_path), width=Inches(doc_width))

                        # Add spacing after image
                        doc.add_paragraph()

                    except Exception as e:
                        logger.error(f"Could not add screenshot {screenshot_path.name} to DOCX: {e}")
                        doc.add_paragraph(f"[Could not load image: {screenshot_path.name}]")
                        doc.add_paragraph()
                else:
                    doc.add_paragraph(f"{idx}. {screenshot_type} - {screenshot_path.name} (File not found)")
                    doc.add_paragraph()

    # Get organized export directory (by week for DOCX)
    exports_dir = get_organized_export_path(db_path, account_name, trade['entry_date'])

    # Save DOCX with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"trade_{trade['trade_id']}_{timestamp}.docx"
    docx_path = exports_dir / filename

    doc.save(str(docx_path))

    return docx_path


# ================================================================================================


st.set_page_config(page_title="Trade History", page_icon="📜", layout="wide", )


icon_header("icons/history.png", "Trade History",width=50, height=50, level="h1")
st.markdown("View, filter, and manage all your trades")
st.markdown("---")

db_path = get_database_path()

# Account selector
selected_account = render_account_selector()

if not selected_account:
    st.stop()

account_id = selected_account['account_id']

# =============================================================================
# ENHANCED FILTERS SECTION
# =============================================================================

icon_header("icons/filter.png", "Filters", level="h3")

# Get all trades for this account to build dynamic filters
all_account_trades = get_all_trades(db_path, account_id)

# Extract unique symbols from trades
unique_symbols = sorted(list(set(t['symbol'] for t in all_account_trades if t.get('symbol'))))

# Get all strategies
all_strategies = get_all_strategies(db_path)

# Build filter UI
col_f1, col_f2, col_f3, col_f4 = st.columns(4)

with col_f1:
    # Symbol filter - dynamic based on actual trades
    symbol_options = ["All Symbols"] + unique_symbols
    selected_symbol = st.selectbox("Symbol", options=symbol_options, key="filter_symbol")

with col_f2:
    # Direction filter
    direction_options = ["All Directions", "LONG", "SHORT"]
    selected_direction = st.selectbox("Direction", options=direction_options, key="filter_direction")

with col_f3:
    # Status filter
    status_options = ["All Status", "OPEN", "CLOSED"]
    selected_status = st.selectbox("Status", options=status_options, key="filter_status")

with col_f4:
    # Strategy filter
    strategy_options = ["All Strategies"] + [s['strategy_name'] for s in all_strategies]
    selected_strategy = st.selectbox("Strategy", options=strategy_options, key="filter_strategy")

# Second row of filters
col_f5, col_f6, col_f7, col_f8 = st.columns(4)

with col_f5:
    # P&L filter
    pnl_options = ["All P&L", "Profit (P&L > 0)", "Loss (P&L < 0)", "Breakeven (P&L ≈ 0)"]
    selected_pnl = st.selectbox("P&L Result", options=pnl_options, key="filter_pnl")

with col_f6:
    # Date from filter
    use_date_from = st.checkbox("Filter by Date Range", key="use_date_filter")

with col_f7:
    date_from = None
    if use_date_from:
        date_from = st.date_input("From Date", key="filter_date_from")

with col_f8:
    date_to = None
    if use_date_from:
        date_to = st.date_input("To Date", key="filter_date_to")

st.markdown("---")

# =============================================================================
# BUILD FILTER DICTIONARY
# =============================================================================

filters = {'account_id': account_id}

# Symbol filter
if selected_symbol != "All Symbols":
    filters['symbol'] = selected_symbol

# Direction filter
if selected_direction != "All Directions":
    filters['direction'] = selected_direction

# Status filter
if selected_status != "All Status":
    filters['status'] = selected_status

# Strategy filter
if selected_strategy != "All Strategies":
    # Find strategy ID
    for s in all_strategies:
        if s['strategy_name'] == selected_strategy:
            filters['strategy_id'] = s['strategy_id']
            break

# Date filters
if use_date_from and date_from:
    filters['date_from'] = date_from.strftime("%Y-%m-%d")

if use_date_from and date_to:
    filters['date_to'] = date_to.strftime("%Y-%m-%d")

# =============================================================================
# GET FILTERED TRADES
# =============================================================================

trades = get_filtered_trades(db_path, filters)

# =============================================================================
# APPLY P&L FILTER (Post-database filter)
# =============================================================================

if selected_pnl != "All P&L":
    filtered_trades = []

    for trade in trades:
        trade_pnl = trade.get('pnl', 0) or 0.0

        if selected_pnl == "Profit (P&L > 0)" and trade_pnl > 0:
            filtered_trades.append(trade)
        elif selected_pnl == "Loss (P&L < 0)" and trade_pnl < 0:
            filtered_trades.append(trade)
        elif selected_pnl == "Breakeven (P&L ≈ 0)" and abs(trade_pnl) < 1.0:
            filtered_trades.append(trade)

    trades = filtered_trades


# =============================================================================
# SORT TRADES BY TRADE_ID DESCENDING (NEWEST FIRST)
# =============================================================================

trades.sort(key=lambda t: t.get('trade_id', 0), reverse=True)

# =============================================================================
# DISPLAY TRADES WITH PAGINATION
# =============================================================================

if trades:
    st.subheader(f"Trades in {selected_account['account_name']}")

    # Summary stats
    col_summary1, col_summary2, col_summary3, col_summary4 = st.columns(4)

    with col_summary1:
        st.metric("Total Trades", len(trades))

    with col_summary2:
        open_count = sum(1 for t in trades if t.get('status') == 'OPEN')
        st.metric("Open Trades", open_count)

    with col_summary3:
        closed_count = sum(1 for t in trades if t.get('status') == 'CLOSED')
        st.metric("Closed Trades", closed_count)

    with col_summary4:
        total_filtered_pnl = sum(t.get('pnl', 0) or 0.0 for t in trades if t.get('status') == 'CLOSED')
        st.metric("Total P&L", format_currency(total_filtered_pnl))

    st.markdown("---")


    # =============================================================================
    # CSV EXPORT (Now saves to CSV_Trades_Exports folder)
    # =============================================================================

    col_exp1, col_exp2, col_exp3 = st.columns([1, 1, 1])

    with col_exp2:
        if st.button("Export to CSV", use_container_width=True):
            csv_rows = []

            for t in trades:
                strategy_name = "No Strategy"
                if t.get('strategy_id'):
                    for s in all_strategies:
                        if s['strategy_id'] == t['strategy_id']:
                            strategy_name = s['strategy_name']
                            break

                csv_rows.append({
                    'Trade ID': t.get('trade_id'),
                    'Symbol': t.get('symbol'),
                    'Direction': t.get('direction'),
                    'Status': t.get('status'),
                    'Entry Date': t.get('entry_date'),
                    'Entry Time': t.get('entry_time'),
                    'Entry Price': t.get('entry_price'),
                    'Position Size': t.get('position_size'),
                    'Stop Loss': t.get('stop_loss'),
                    'Target': t.get('target'),
                    'Risk Amount': t.get('risk_amount'),
                    'Risk %': t.get('risk_percentage'),
                    'FTA': t.get('fta'),
                    'Prospective R': t.get('prospective_r'),
                    'Trading Environment': t.get('trading_environment'),
                    'Strategy': strategy_name,
                    'Trigger': t.get('trigger'),
                    'Exit Date': t.get('exit_date'),
                    'Exit Time': t.get('exit_time'),
                    'Exit Price': t.get('multi_exit_price'),
                    'P&L': t.get('pnl'),
                    'Total R': t.get('total_r'),
                    'ROE %': t.get('roe_percentage'),
                    'MAE': t.get('mae'),
                    'MFE': t.get('mfe'),
                    'Duration': t.get('trade_duration'),
                    'Target Hit': t.get('target_hit'),
                    'Stop Loss Hit': t.get('stop_loss_hit'),
                    'Grade Mentally': t.get('grade_mentally'),
                    'Grade Technically': t.get('grade_technically'),
                    'Setup Idea': t.get('setup_idea'),
                    'Notes Entry': t.get('trade_notes_entry'),
                    'Notes Management': t.get('trade_notes_management'),
                    'Notes Closing': t.get('trade_notes_closing'),
                    'Reason Closing': t.get('reason_for_closing'),
                    'Final Notes': t.get('final_notes')
                })

            df = pd.DataFrame(csv_rows)

            # Get CSV exports directory
            csv_exports_dir = get_csv_exports_dir(db_path)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"trades_{selected_account['account_name']}_{timestamp}.csv"

            # Save CSV file
            csv_path = csv_exports_dir / filename
            df.to_csv(csv_path, index=False)

            st.success(f"CSV exported successfully!")
            st.info(f"File saved to: {csv_path}")
            st.caption(f"Total trades exported: {len(trades)}")

    st.markdown("---")

    # =============================================================================
    # PAGINATION LOGIC
    # =============================================================================

    trades_per_page = 50
    total_trades = len(trades)
    total_pages = (total_trades + trades_per_page - 1) // trades_per_page

    # Show pagination controls if more than 50 trades
    if total_pages > 1:
        col_page1, col_page2, col_page3 = st.columns([1, 2, 1])
        with col_page2:
            page = st.number_input(
                f"🗎 Page (1-{total_pages})", 
                min_value=1, 
                max_value=total_pages, 
                value=1,
                key="trade_history_page"
            )
    else:
        page = 1

    # Calculate slice indices
    start_idx = (page - 1) * trades_per_page
    end_idx = start_idx + trades_per_page

    # Display only current page
    displayed_trades = trades[start_idx:end_idx]

    # Show pagination info
    st.info(f"Showing trades {start_idx + 1}-{min(end_idx, total_trades)} of {total_trades} total")
    st.caption("Click on a trade row to view full details")
    st.markdown("---")

    # =============================================================================
    # TRADE TABLE VIEW
    # =============================================================================

    table_rows = []
    for t in displayed_trades:
        strategy_name = "—"
        if t.get('strategy_id'):
            for s in all_strategies:
                if s['strategy_id'] == t['strategy_id']:
                    strategy_name = s['strategy_name']
                    break

        pnl_val = t.get('pnl', 0) or 0.0
        r_val   = t.get('total_r')
        status  = t.get('status', '')

        table_rows.append({
            'ID':        f"#{t['trade_id']}",
            'Symbol':    t.get('symbol', ''),
            'Direction': t.get('direction', ''),
            'Status':    status,
            'Date':      t.get('entry_date', ''),
            'Entry $':   f"{t.get('entry_price', 0):.4f}",
            'SL $':      f"{t.get('stop_loss', 0):.4f}",
            'Size':      f"{t.get('position_size', 0):.4f}",
            'R':         f"{r_val:.2f}R" if r_val is not None else "—",
            'P&L':       f"${pnl_val:+,.2f}" if status == 'CLOSED' else "OPEN",
            'Strategy':  strategy_name,
        })

    df_display = pd.DataFrame(table_rows)
    st.dataframe(
        df_display,
        use_container_width=True,
        hide_index=True,
        height=min(50 + len(table_rows) * 35, 600)
    )

    st.markdown("---")

    # =============================================================================
    # TRADE DETAIL / EDIT PANEL
    # =============================================================================

    trade_id_options = [
        f"#{t['trade_id']} — {t['symbol']} {t['direction']} ({t['status']})"
        for t in displayed_trades
    ]

    selected_trade_label = st.selectbox(
        "Select trade to view details or edit",
        options=["— select —"] + trade_id_options,
        key="detail_trade_select"
    )

    if selected_trade_label != "— select —":
        sel_idx = trade_id_options.index(selected_trade_label)
        trade   = displayed_trades[sel_idx]

        tab_detail, tab_notes, tab_screenshots, tab_export, tab_danger = st.tabs([
            "Details", "Notes", "Screenshots", "Export", "Delete"
        ])

        # =====================================================================
        # TAB: DETAILS
        # =====================================================================

        with tab_detail:

            col_detail1, col_detail2 = st.columns(2)

            with col_detail1:
                icon_header("icons/trade_details.png", "Trade Details", level="h3")
                st.write(f"**Trade ID:** #{trade['trade_id']}")
                st.write(f"**Symbol:** {trade['symbol']}")
                st.write(f"**Direction:** {format_direction(trade['direction'])}")
                st.write(f"**Status:** {trade['status']}")
                st.write(f"**Entry Date:** {trade['entry_date']}")
                st.write(f"**Entry Time:** {trade['entry_time']}")
                st.write(f"**Entry Price:** {format_currency(trade['entry_price'])}")

                if trade.get('multi_entry_price'):
                    st.write(f"**Weighted Avg Entry:** {format_currency(trade['multi_entry_price'])}")

                st.markdown("---")

                icon_header("icons/trade_setup.png", "Trade Setup", level="h3")

                trading_env = trade.get('trading_environment')
                st.write(f"**Trading Environment:** {trading_env if trading_env else 'N/A'}")

                strategy_name = "No Strategy"
                if trade.get('strategy_id'):
                    for s in all_strategies:
                        if s['strategy_id'] == trade['strategy_id']:
                            strategy_name = s['strategy_name']
                            break
                st.write(f"**Strategy:** {strategy_name}")

                trigger = trade.get('trigger')
                st.write(f"**Trigger:** {trigger if trigger else 'N/A'}")

            with col_detail2:
                icon_header("icons/balance.png", "Position Sizing & Risk", level="h3")
                st.write(f"**Position Size:** {trade['position_size']:.8f} units")
                st.write(f"**Stop Loss:** {format_currency(trade['stop_loss'])}")

                target = trade.get('target')
                st.write(f"**Target:** {format_currency(target) if target else 'N/A'}")

                st.write(f"**Risk Amount:** {format_currency(trade['risk_amount'])}")
                st.write(f"**Risk %:** {format_percentage(trade['risk_percentage'])}")

                fta = trade.get('fta')
                st.write(f"**FTA:** {format_currency(fta) if fta else 'N/A'}")

                prosp_r = trade.get('prospective_r')
                st.write(f"**Prospective R:** {format_r_multiple(prosp_r) if prosp_r else 'N/A'}")

                within_risk = trade.get('within_risk_limit', True)
                st.write(f"**Within Risk Limit:** {'[ OK ] Yes' if within_risk else '[ X ] No'}")

                comm = trade.get('commission')
                st.write(f"**Commission:** {format_currency(comm) if comm else 'N/A'}")

            # =================================================================
            # TRADE RESULTS (IF CLOSED)
            # =================================================================

            if trade['status'] == 'CLOSED':
                st.markdown("---")
                icon_header("icons/summarize.png", "Trade Results", level="h3")

                col_res1, col_res2, col_res3 = st.columns(3)

                with col_res1:
                    st.write(f"**Exit Date:** {trade.get('exit_date', 'N/A')}")
                    st.write(f"**Exit Time:** {trade.get('exit_time', 'N/A')}")

                    exit_price = trade.get('multi_exit_price')
                    st.write(f"**Exit Price:** {format_currency(exit_price) if exit_price else 'N/A'}")

                    target_hit = trade.get('target_hit', False)
                    st.write(f"**Target Hit:** {'[ OK ] Yes' if target_hit else '[ X ] No'}")

                    stop_hit = trade.get('stop_loss_hit', False)
                    st.write(f"**Stop Loss Hit:** {'[ OK ] Yes' if stop_hit else '[ X ] No'}")

                with col_res2:
                    mae = trade.get('mae')
                    st.write(f"**MAE:** {format_currency(mae) if mae else 'N/A'}")

                    mfe = trade.get('mfe')
                    st.write(f"**MFE:** {format_currency(mfe) if mfe else 'N/A'}")

                    duration = trade.get('trade_duration')
                    st.write(f"**Duration:** {format_duration(duration) if duration else 'N/A'}")

                    total_r = trade.get('total_r')
                    st.write(f"**Total R:** {format_r_multiple(total_r) if total_r else 'N/A'}")

                    roe = trade.get('roe_percentage')
                    st.write(f"**ROE %:** {format_percentage(roe) if roe is not None else 'N/A'}")

                with col_res3:
                    trade_pnl = trade.get('pnl', 0) or 0.0

                    if trade_pnl > 0:
                        st.success(f"**P&L:** {format_currency(trade_pnl)}")
                    elif trade_pnl < 0:
                        st.error(f"**P&L:** {format_currency(trade_pnl)}")
                    else:
                        st.info(f"**P&L:** {format_currency(trade_pnl)}")

                st.markdown("---")
                icon_header("icons/doc_list.png", "Trade Grades", level="h3")

                col_grade1, col_grade2 = st.columns(2)

                with col_grade1:
                    grade_mental = trade.get('grade_mentally')
                    st.write(f"**Grade Mentally:** {grade_mental if grade_mental else 'Not Graded'}")

                with col_grade2:
                    grade_tech = trade.get('grade_technically')
                    st.write(f"**Grade Technically:** {grade_tech if grade_tech else 'Not Graded'}")

            # =================================================================
            # EDIT PANEL — OPEN TRADES ONLY
            # =================================================================

            if trade['status'] == 'OPEN':
                st.markdown("---")
                icon_header("icons/trade_details.png", "Edit Open Trade", level="h3")
                st.caption("Only open trades can be edited. To close a trade go to Log Trade -> Close Existing Trade.")

                with st.form(key=f"edit_open_trade_{trade['trade_id']}"):
                    ec1, ec2 = st.columns(2)

                    with ec1:
                        ed_symbol = st.text_input(
                            "Symbol",
                            value=trade.get('symbol', ''),
                            key=f"ed_sym_{trade['trade_id']}"
                        ).upper()

                        ed_entry_price = st.number_input(
                            "Entry Price",
                            value=float(trade.get('entry_price') or 0),
                            min_value=0.0,
                            format="%.8f",
                            key=f"ed_ep_{trade['trade_id']}"
                        )

                        ed_stop_loss = st.number_input(
                            "Stop Loss",
                            value=float(trade.get('stop_loss') or 0),
                            min_value=0.0,
                            format="%.8f",
                            key=f"ed_sl_{trade['trade_id']}"
                        )

                        ed_target = st.number_input(
                            "Target (0 = none)",
                            value=float(trade.get('target') or 0),
                            min_value=0.0,
                            format="%.8f",
                            key=f"ed_tgt_{trade['trade_id']}"
                        )

                        ed_position_size = st.number_input(
                            "Position Size",
                            value=float(trade.get('position_size') or 0),
                            min_value=0.0,
                            format="%.8f",
                            key=f"ed_ps_{trade['trade_id']}"
                        )

                    with ec2:
                        ed_risk_amount = st.number_input(
                            "Risk Amount ($)",
                            value=float(trade.get('risk_amount') or 0),
                            min_value=0.0,
                            format="%.2f",
                            key=f"ed_ra_{trade['trade_id']}"
                        )

                        ed_risk_pct = st.number_input(
                            "Risk %",
                            value=float(trade.get('risk_percentage') or 0),
                            min_value=0.0,
                            max_value=100.0,
                            format="%.4f",
                            key=f"ed_rp_{trade['trade_id']}"
                        )

                        ed_commission = st.number_input(
                            "Commission",
                            value=float(trade.get('commission') or 0),
                            min_value=0.0,
                            format="%.8f",
                            key=f"ed_com_{trade['trade_id']}"
                        )

                        ed_trigger = st.text_input(
                            "Trigger",
                            value=trade.get('trigger') or '',
                            key=f"ed_trg_{trade['trade_id']}"
                        )

                        from config.constants import TRADING_ENVIRONMENTS
                        _EXTRA_ENVS = ["Trending UP", "Trending DOWN", "Trend Shift"]
                        ed_env_options = [""] + TRADING_ENVIRONMENTS + [
                            e for e in _EXTRA_ENVS if e not in TRADING_ENVIRONMENTS
                        ]
                        current_env = trade.get('trading_environment') or ""
                        ed_env_idx  = ed_env_options.index(current_env) if current_env in ed_env_options else 0

                        ed_env = st.selectbox(
                            "Trading Environment",
                            options=ed_env_options,
                            index=ed_env_idx,
                            key=f"ed_env_{trade['trade_id']}"
                        )

                    ed_within_risk = st.checkbox(
                        "Within Risk Limit",
                        value=bool(trade.get('within_risk_limit', True)),
                        key=f"ed_wr_{trade['trade_id']}"
                    )

                    ed_setup_idea = st.text_area(
                        "Setup Idea",
                        value=trade.get('setup_idea') or '',
                        key=f"ed_si_{trade['trade_id']}"
                    )

                    ed_notes_entry = st.text_area(
                        "Trade Notes (Entry)",
                        value=trade.get('trade_notes_entry') or '',
                        key=f"ed_ne_{trade['trade_id']}"
                    )

                    save_btn = st.form_submit_button(
                        "Save Changes",
                        type="primary",
                        use_container_width=True
                    )

                if save_btn:
                    updated_fields = {
                        'symbol':              ed_symbol,
                        'entry_price':         ed_entry_price,
                        'multi_entry_price':   ed_entry_price,
                        'stop_loss':           ed_stop_loss,
                        'target':              ed_target if ed_target > 0 else None,
                        'position_size':       ed_position_size,
                        'risk_amount':         ed_risk_amount,
                        'risk_percentage':     ed_risk_pct,
                        'commission':          ed_commission if ed_commission > 0 else None,
                        'trigger':             ed_trigger if ed_trigger else None,
                        'trading_environment': ed_env if ed_env else None,
                        'within_risk_limit':   ed_within_risk,
                        'setup_idea':          ed_setup_idea if ed_setup_idea else None,
                        'trade_notes_entry':   ed_notes_entry if ed_notes_entry else None,
                    }

                    success = update_trade_open_fields(db_path, trade['trade_id'], updated_fields)

                    if success:
                        st.success(f"[ OK ] Trade #{trade['trade_id']} updated successfully!")
                        clear_cache_after_trade_operation('edit')
                        st.rerun()
                    else:
                        st.error("[ X ] Failed to save changes.")

        # =====================================================================
        # TAB: NOTES
        # =====================================================================

        with tab_notes:
            icon_header("icons/doc_list.png", "Trade Notes", level="h3")

            has_notes = False

            if trade.get('setup_idea'):
                st.markdown("**Setup Idea:**")
                st.info(trade['setup_idea'])
                has_notes = True

            if trade.get('trade_notes_entry'):
                st.markdown("**Trade Notes (Entry):**")
                st.info(trade['trade_notes_entry'])
                has_notes = True

            if trade.get('trade_notes_management'):
                st.markdown("**Trade Notes (Management):**")
                st.info(trade['trade_notes_management'])
                has_notes = True

            if trade.get('trade_notes_closing'):
                st.markdown("**Trade Notes (Closing):**")
                st.info(trade['trade_notes_closing'])
                has_notes = True

            if trade.get('reason_for_closing'):
                st.markdown("**Reason for Closing:**")
                st.warning(trade['reason_for_closing'])
                has_notes = True

            if trade.get('final_notes'):
                st.markdown("**Final Notes:**")
                st.info(trade['final_notes'])
                has_notes = True

            if not has_notes:
                st.caption("No notes available for this trade.")

        # =====================================================================
        # TAB: SCREENSHOTS
        # =====================================================================

        with tab_screenshots:
            icon_header("icons/doc_list.png", "Screenshots", level="h3")

            screenshots = get_trade_screenshots(db_path, trade['trade_id'])

            if screenshots and len(screenshots) > 0:
                st.success(f"Found {len(screenshots)} screenshot(s) for this trade")

                entry_screenshots = [s for s in screenshots if s.get('screenshot_type') == 'ENTRY']
                exit_screenshots  = [s for s in screenshots if s.get('screenshot_type') == 'EXIT']
                other_screenshots = [s for s in screenshots if s.get('screenshot_type') not in ('ENTRY', 'EXIT')]

                for group_label, group in [
                    ("Entry Screenshots",  entry_screenshots),
                    ("Exit Screenshots",   exit_screenshots),
                    ("Other Screenshots",  other_screenshots),
                ]:
                    if not group:
                        continue

                    st.markdown(f"#### {group_label}")
                    cols = st.columns(min(len(group), 3))

                    for idx, screenshot in enumerate(group):
                        with cols[idx % 3]:
                            try:
                                file_path = screenshot.get('file_path', '')

                                if not file_path:
                                    st.error("[ X ] No file path in database")
                                    continue

                                screenshot_path = Path(file_path)
                                st.caption(f"**Path:** `{str(screenshot_path)}`")

                                if screenshot_path.exists():
                                    st.success("[ OK ] File exists")
                                    try:
                                        image = Image.open(screenshot_path)
                                        st.image(
                                            image,
                                            caption=screenshot_path.name,
                                            use_container_width=True
                                        )
                                    except Exception as img_error:
                                        st.error(f"[ X ] Cannot open image: {str(img_error)}")
                                else:
                                    st.error("[ X ] File not found at path")
                                    st.caption(f"Expected: `{str(screenshot_path.absolute())}`")

                            except Exception as e:
                                st.error(f"[ X ] Error: {str(e)}")
                                logger.error(f"Screenshot error for trade {trade['trade_id']}: {e}")

                    st.markdown("---")

                if not entry_screenshots and not exit_screenshots and not other_screenshots:
                    st.warning("[ ! ] Screenshots found in database but no valid types detected")
            else:
                st.info("[ ! ] No screenshots attached to this trade yet")

        # =====================================================================
        # TAB: EXPORT
        # =====================================================================

        with tab_export:
            icon_header("icons/doc_list.png", "Export Trade", level="h3")

            col_pdf1, col_pdf2, col_pdf3 = st.columns([2, 1, 2])

            with col_pdf2:
                if st.button(
                    "Export to Word",
                    key=f"export_docx_{trade['trade_id']}",
                    use_container_width=True
                ):
                    try:
                        screenshots = get_trade_screenshots(db_path, trade['trade_id'])
                        docx_path = export_trade_to_docx(
                            trade,
                            selected_account['account_name'],
                            all_strategies,
                            screenshots,
                            db_path
                        )
                        st.success("[ OK ] Word document exported successfully!")
                        st.info(f"File saved to: {docx_path}")

                        import subprocess
                        try:
                            subprocess.Popen(f'explorer /select,"{str(docx_path)}"')
                            st.success("[ OK ] Opening file location in Explorer...")
                        except Exception as explorer_error:
                            st.warning(f"[ ! ] Could not open Explorer automatically: {explorer_error}")
                            st.info(f"Please navigate to: {docx_path.parent}")

                    except Exception as e:
                        st.error(f"[ X ] Failed to export document: {str(e)}")
                        import traceback
                        st.code(traceback.format_exc())

        # =====================================================================
        # TAB: DELETE
        # =====================================================================

        with tab_danger:
            icon_header("icons/danger.png", "Danger Zone", level="h3")

            col_del1, col_del2, col_del3 = st.columns([2, 1, 2])

            with col_del2:
                delete_confirm = st.checkbox(
                    "Confirm Delete",
                    key=f"confirm_delete_{trade['trade_id']}",
                    help="Check this box to enable the delete button"
                )

                if st.button(
                    "Delete Trade",
                    type="secondary",
                    use_container_width=True,
                    disabled=not delete_confirm,
                    key=f"delete_btn_{trade['trade_id']}"
                ):
                    success = delete_trade(db_path, trade['trade_id'])

                    if success:
                        st.success(f"[ OK ] Trade #{trade['trade_id']} deleted successfully!")
                        clear_cache_after_trade_operation('delete')
                        st.rerun()
                    else:
                        st.error(f"[ X ] Failed to delete trade #{trade['trade_id']}")

else:
    st.info("[ ! ] No trades found matching the selected filters.")

st.markdown("---")
st.caption("💡 **Tips:**")
st.caption("• Use filters to find specific trades quickly")
st.caption("• P&L filter includes: Profit (>$0), Loss (<$0), Breakeven (±$1)")
st.caption("• Click 'Confirm Delete' checkbox before deleting a trade")
st.caption("• All screenshots are organized by type (Entry, Exit, Other)")
